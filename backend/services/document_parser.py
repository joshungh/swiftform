import os
import re
from typing import Dict, List, Any, Optional
import PyPDF2
import pdfplumber
from docx import Document
import openpyxl
import pandas as pd
from PIL import Image
import pytesseract
import json

class DocumentParser:
    """Service for parsing various document formats"""

    def __init__(self):
        self.supported_formats = {'.pdf', '.doc', '.docx', '.xls', '.xlsx'}

    async def parse_document(self, file_path: str) -> Dict[str, Any]:
        """Main entry point for document parsing"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_ext = os.path.splitext(file_path)[1].lower()

        if file_ext == '.pdf':
            return await self.parse_pdf(file_path)
        elif file_ext in ['.doc', '.docx']:
            return await self.parse_word(file_path)
        elif file_ext in ['.xls', '.xlsx']:
            return await self.parse_excel(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

    async def parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Parse PDF documents"""
        content = {
            "type": "pdf",
            "pages": [],
            "forms": [],
            "tables": [],
            "text_blocks": [],
            "metadata": {}
        }

        try:
            with pdfplumber.open(file_path) as pdf:
                content["metadata"] = {
                    "pages": len(pdf.pages),
                    "author": pdf.metadata.get('Author', ''),
                    "title": pdf.metadata.get('Title', ''),
                    "subject": pdf.metadata.get('Subject', '')
                }

                for page_num, page in enumerate(pdf.pages, 1):
                    page_data = {
                        "page_number": page_num,
                        "text": page.extract_text() or "",
                        "tables": [],
                        "form_fields": []
                    }

                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            page_data["tables"].append({
                                "data": table,
                                "headers": table[0] if table else []
                            })

                    page_data["form_fields"] = self._extract_form_fields(page_data["text"])

                    content["pages"].append(page_data)

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                if pdf_reader.is_encrypted:
                    pdf_reader.decrypt('')

                if '/AcroForm' in pdf_reader.trailer['/Root']:
                    content["forms"] = self._extract_pdf_forms(pdf_reader)

        except Exception as e:
            content["error"] = str(e)

        return content

    async def parse_word(self, file_path: str) -> Dict[str, Any]:
        """Parse Word documents"""
        content = {
            "type": "word",
            "paragraphs": [],
            "tables": [],
            "headers": [],
            "form_fields": [],
            "metadata": {}
        }

        try:
            doc = Document(file_path)

            content["metadata"] = {
                "author": doc.core_properties.author or "",
                "title": doc.core_properties.title or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else ""
            }

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    para_data = {
                        "text": paragraph.text,
                        "style": paragraph.style.name if paragraph.style else "",
                        "level": self._get_heading_level(paragraph.style.name if paragraph.style else "")
                    }
                    content["paragraphs"].append(para_data)

                    if para_data["level"] > 0:
                        content["headers"].append(para_data)

            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                content["tables"].append({
                    "data": table_data,
                    "headers": table_data[0] if table_data else []
                })

            full_text = "\n".join([p["text"] for p in content["paragraphs"]])
            content["form_fields"] = self._extract_form_fields(full_text)

        except Exception as e:
            content["error"] = str(e)

        return content

    async def parse_excel(self, file_path: str) -> Dict[str, Any]:
        """Parse Excel documents"""
        content = {
            "type": "excel",
            "sheets": [],
            "form_fields": [],
            "metadata": {}
        }

        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            content["metadata"]["sheet_names"] = workbook.sheetnames

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_data = {
                    "name": sheet_name,
                    "data": [],
                    "headers": [],
                    "form_fields": []
                }

                df = pd.read_excel(file_path, sheet_name=sheet_name)
                sheet_data["data"] = df.to_dict(orient='records')
                sheet_data["headers"] = df.columns.tolist()

                text_content = df.to_string()
                sheet_data["form_fields"] = self._extract_form_fields(text_content)

                content["sheets"].append(sheet_data)

        except Exception as e:
            content["error"] = str(e)

        return content

    def _extract_form_fields(self, text: str) -> List[Dict[str, Any]]:
        """Extract potential form fields from text"""
        fields = []

        patterns = {
            'date': r'\b(?:date|Date|DATE)\s*:?\s*(?:_+|\[.*?\]|\(.*?\))?',
            'name': r'\b(?:name|Name|NAME)\s*:?\s*(?:_+|\[.*?\]|\(.*?\))?',
            'email': r'\b(?:email|Email|EMAIL)\s*:?\s*(?:_+|\[.*?\]|\(.*?\))?',
            'phone': r'\b(?:phone|Phone|PHONE|tel|Tel|TEL)\s*:?\s*(?:_+|\[.*?\]|\(.*?\))?',
            'address': r'\b(?:address|Address|ADDRESS)\s*:?\s*(?:_+|\[.*?\]|\(.*?\))?',
            'checkbox': r'[\[\]☐☑✓✗]\s*(.+?)(?:\n|$)',
            'radio': r'[○●◯◉]\s*(.+?)(?:\n|$)',
            'select': r'\b(?:select|Select|SELECT|choose|Choose|CHOOSE)\s*:?\s*(?:_+|\[.*?\]|\(.*?\))?',
            'text_field': r'(?:^|\n)([^:]+?):\s*_+',
            'label_field': r'(?:^|\n)([^:]+?):\s*(?:\[.*?\]|\(.*?\))',
            'number': r'\b(?:number|Number|NUMBER|amount|Amount|AMOUNT|quantity|Quantity|QUANTITY)\s*:?\s*(?:_+|\[.*?\]|\(.*?\))?',
        }

        for field_type, pattern in patterns.items():
            matches = re.finditer(pattern, text, re.MULTILINE)
            for match in matches:
                field_text = match.group(1) if match.groups() else match.group(0)
                field_text = field_text.strip()

                if field_text and len(field_text) < 100:
                    fields.append({
                        "type": field_type,
                        "label": field_text,
                        "original_text": match.group(0),
                        "position": match.start()
                    })

        fields.sort(key=lambda x: x["position"])

        return fields

    def _extract_pdf_forms(self, pdf_reader) -> List[Dict[str, Any]]:
        """Extract form fields from PDF AcroForms"""
        forms = []

        try:
            if '/AcroForm' in pdf_reader.trailer['/Root']:
                acro_form = pdf_reader.trailer['/Root']['/AcroForm']
                if '/Fields' in acro_form:
                    for field in acro_form['/Fields']:
                        field_obj = field.get_object()
                        form_field = {
                            "name": field_obj.get('/T', ''),
                            "type": field_obj.get('/FT', ''),
                            "value": field_obj.get('/V', ''),
                            "default": field_obj.get('/DV', ''),
                            "flags": field_obj.get('/Ff', 0)
                        }
                        forms.append(form_field)
        except Exception:
            pass

        return forms

    def _get_heading_level(self, style_name: str) -> int:
        """Determine heading level from style name"""
        if not style_name:
            return 0

        style_lower = style_name.lower()
        if 'heading 1' in style_lower:
            return 1
        elif 'heading 2' in style_lower:
            return 2
        elif 'heading 3' in style_lower:
            return 3
        elif 'heading' in style_lower:
            return 4
        elif 'title' in style_lower:
            return 1

        return 0

    def extract_structure(self, content: Dict[str, Any]) -> Dict[str, Any]:
        """Extract document structure for AI processing"""
        structure = {
            "sections": [],
            "fields": [],
            "tables": [],
            "lists": []
        }

        if content["type"] == "pdf":
            for page in content.get("pages", []):
                structure["fields"].extend(page.get("form_fields", []))
                structure["tables"].extend(page.get("tables", []))

                text = page.get("text", "")
                sections = self._extract_sections(text)
                structure["sections"].extend(sections)

        elif content["type"] == "word":
            structure["fields"] = content.get("form_fields", [])
            structure["tables"] = content.get("tables", [])

            current_section = None
            for para in content.get("paragraphs", []):
                if para.get("level", 0) > 0:
                    if current_section:
                        structure["sections"].append(current_section)
                    current_section = {
                        "title": para["text"],
                        "level": para["level"],
                        "content": []
                    }
                elif current_section:
                    current_section["content"].append(para["text"])

            if current_section:
                structure["sections"].append(current_section)

        elif content["type"] == "excel":
            for sheet in content.get("sheets", []):
                structure["fields"].extend(sheet.get("form_fields", []))
                structure["tables"].append({
                    "name": sheet["name"],
                    "headers": sheet.get("headers", []),
                    "data": sheet.get("data", [])
                })

        return structure

    def _extract_sections(self, text: str) -> List[Dict[str, Any]]:
        """Extract sections from text based on patterns"""
        sections = []

        section_patterns = [
            r'^#+\s+(.+)$',
            r'^(\d+\.?\s+[A-Z].+)$',
            r'^([A-Z][A-Z\s]+):?\s*$',
            r'^([IVX]+\.?\s+.+)$',
        ]

        lines = text.split('\n')
        current_section = None

        for line in lines:
            line = line.strip()
            if not line:
                continue

            is_section = False
            for pattern in section_patterns:
                match = re.match(pattern, line)
                if match:
                    if current_section:
                        sections.append(current_section)
                    current_section = {
                        "title": match.group(1).strip(),
                        "content": []
                    }
                    is_section = True
                    break

            if not is_section and current_section:
                current_section["content"].append(line)

        if current_section:
            sections.append(current_section)

        return sections